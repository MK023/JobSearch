"""Tests for document scanning service."""

import io
import json
from unittest.mock import MagicMock, patch

from src.integrations.document_scanner import (
    _extract_text_from_docx,
    scan_document,
)
from src.interview.file_models import FileStatus


class TestExtractTextFromDocx:
    def test_extracts_paragraph_text(self):
        """Create a minimal DOCX in memory and extract text."""
        from docx import Document

        doc = Document()
        doc.add_paragraph("Nome: Mario Rossi")
        doc.add_paragraph("Indirizzo: Via Roma 1")

        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        text = _extract_text_from_docx(buffer.read())
        assert "Mario Rossi" in text
        assert "Via Roma 1" in text

    def test_empty_docx(self):
        from docx import Document

        doc = Document()
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        text = _extract_text_from_docx(buffer.read())
        assert text.strip() == ""


class TestScanDocument:
    @patch("src.integrations.document_scanner.get_client")
    def test_scan_pdf_compiled(self, mock_get_client):
        mock_client = MagicMock()
        mock_message = MagicMock()
        mock_message.content = [
            MagicMock(
                text=json.dumps(
                    {
                        "compiled": True,
                        "confidence": "high",
                        "summary": "Documento compilato con dati personali.",
                    }
                )
            )
        ]
        mock_message.usage = MagicMock(input_tokens=500, output_tokens=50)
        mock_client.messages.create.return_value = mock_message
        mock_get_client.return_value = mock_client

        result = scan_document(
            file_bytes=b"%PDF-1.4 fake content",
            filename="modulo.pdf",
            content_type="application/pdf",
        )

        assert result["status"] == FileStatus.COMPILED
        assert result["compiled"] is True
        assert result["confidence"] == "high"
        assert result["cost_usd"] > 0

    @patch("src.integrations.document_scanner.get_client")
    def test_scan_pdf_not_compiled(self, mock_get_client):
        mock_client = MagicMock()
        mock_message = MagicMock()
        mock_message.content = [
            MagicMock(
                text=json.dumps(
                    {
                        "compiled": False,
                        "confidence": "high",
                        "summary": "Template vuoto con campi da compilare.",
                    }
                )
            )
        ]
        mock_message.usage = MagicMock(input_tokens=500, output_tokens=50)
        mock_client.messages.create.return_value = mock_message
        mock_get_client.return_value = mock_client

        result = scan_document(
            file_bytes=b"%PDF-1.4 fake content",
            filename="template.pdf",
            content_type="application/pdf",
        )

        assert result["status"] == FileStatus.NOT_COMPILED
        assert result["compiled"] is False

    def test_scan_empty_docx(self):
        """Empty DOCX returns NOT_COMPILED without calling API."""
        from docx import Document

        doc = Document()
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        result = scan_document(
            file_bytes=buffer.read(),
            filename="empty.docx",
            content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )

        assert result["status"] == FileStatus.NOT_COMPILED
        assert result["compiled"] is False
        assert result["confidence"] == "high"
        assert result["cost_usd"] == 0.0

    @patch("src.integrations.document_scanner.get_client")
    def test_scan_error_returns_scan_error(self, mock_get_client):
        mock_client = MagicMock()
        mock_client.messages.create.side_effect = Exception("API error")
        mock_get_client.return_value = mock_client

        result = scan_document(
            file_bytes=b"%PDF-1.4 fake content",
            filename="broken.pdf",
            content_type="application/pdf",
        )

        assert result["status"] == FileStatus.SCAN_ERROR
        assert result["compiled"] is False

    @patch("src.integrations.document_scanner.get_client")
    def test_scan_docx_with_content(self, mock_get_client):
        """DOCX with content calls Claude API."""
        from docx import Document

        doc = Document()
        doc.add_paragraph("Nome: Mario Rossi")
        doc.add_paragraph("Cognome: Rossi")
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        mock_client = MagicMock()
        mock_message = MagicMock()
        mock_message.content = [
            MagicMock(
                text=json.dumps(
                    {
                        "compiled": True,
                        "confidence": "high",
                        "summary": "Documento compilato.",
                    }
                )
            )
        ]
        mock_message.usage = MagicMock(input_tokens=300, output_tokens=40)
        mock_client.messages.create.return_value = mock_message
        mock_get_client.return_value = mock_client

        result = scan_document(
            file_bytes=buffer.read(),
            filename="filled.docx",
            content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )

        assert result["status"] == FileStatus.COMPILED
        assert result["compiled"] is True
        mock_client.messages.create.assert_called_once()
