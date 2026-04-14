"""Document scanning service using Claude API.

Analyzes uploaded PDF/DOCX/XLSX/TXT files to determine if they have been
properly filled/compiled by the user, or are still blank/template.
"""

import base64
import io
import logging
from typing import Any

import anthropic
from docx import Document as DocxDocument
from openpyxl import load_workbook

from ..interview.file_models import FileStatus
from .anthropic_client import MODELS, _calculate_cost, get_client

logger = logging.getLogger(__name__)

SCAN_SYSTEM_PROMPT = """\
Sei un assistente che analizza documenti relativi a candidature lavorative.
Il tuo compito e' determinare se il documento e' stato compilato/riempito dall'utente
oppure se e' ancora un modello vuoto/template non compilato.

Criteri per "compiled: true":
- Il documento contiene dati personali specifici (nome, cognome, indirizzo, ecc.)
- I campi del modulo sono stati riempiti con informazioni reali
- Il documento ha contenuto sostanziale oltre alle intestazioni/template

Criteri per "compiled: false":
- Il documento e' un template vuoto con campi placeholder (es. "___", "[Nome]", "INSERIRE QUI")
- Il documento contiene solo intestazioni senza contenuto
- Il documento e' praticamente vuoto
"""

SCAN_TOOL_NAME = "submit_scan_result"
SCAN_TOOL_DESCRIPTION = "Emit the structured document scan result."
SCAN_INPUT_SCHEMA: dict[str, Any] = {
    "type": "object",
    "properties": {
        "compiled": {"type": "boolean", "description": "True if the document has been filled out."},
        "confidence": {"type": "string", "enum": ["high", "medium", "low"]},
        "summary": {"type": "string", "description": "Short description, max 200 chars."},
    },
    "required": ["compiled", "confidence", "summary"],
}

SCAN_USER_PROMPT = """\
Analizza il seguente documento e determina se e' stato compilato o meno.

Nome file: {filename}
Tipo: {content_type}

Contenuto del documento:
{content}
"""


def _extract_text_from_docx(file_bytes: bytes) -> str:
    """Extract text content from a DOCX file."""
    doc = DocxDocument(io.BytesIO(file_bytes))
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    # Also extract from tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text = cell.text.strip()
                if text:
                    paragraphs.append(text)
    return "\n".join(paragraphs)


def _extract_text_from_xlsx(file_bytes: bytes) -> str:
    """Extract text content from an XLSX file."""
    wb = load_workbook(io.BytesIO(file_bytes), read_only=True, data_only=True)
    lines: list[str] = []
    for sheet in wb.worksheets:
        lines.append(f"[Foglio: {sheet.title}]")
        for row in sheet.iter_rows(values_only=True):
            cells = [str(c) for c in row if c is not None]
            if cells:
                lines.append(" | ".join(cells))
    wb.close()
    return "\n".join(lines)


def scan_document(
    file_bytes: bytes,
    filename: str,
    content_type: str,
    model: str = "haiku",
) -> dict[str, Any]:
    """Scan a document with Claude API to check if it's been compiled.

    For PDFs: sends as base64 document to Claude's vision/document support.
    For DOCX: extracts text first, then sends as text.

    Args:
        file_bytes: Raw file content.
        filename: Original filename.
        content_type: MIME type.
        model: Claude model to use (default: haiku for cost savings).

    Returns:
        dict with keys:
        - status: FileStatus.COMPILED or FileStatus.NOT_COMPILED or FileStatus.SCAN_ERROR
        - scan_result: summary string from Claude
        - compiled: bool
        - confidence: str
        - cost_usd: float
        - tokens: dict
    """
    model_id = MODELS.get(model, MODELS["haiku"])
    client = get_client()

    # Content types that need XLSX extraction
    XLSX_TYPES = {
        "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        "application/vnd.ms-excel",
    }
    # Content types that need DOCX extraction
    DOCX_TYPES = {
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        "application/msword",
    }

    try:
        if content_type == "application/pdf":
            return _scan_pdf(client, file_bytes, filename, content_type, model_id)
        if content_type in XLSX_TYPES:
            return _scan_text_based(
                client,
                _extract_text_from_xlsx(file_bytes),
                filename,
                content_type,
                model_id,
            )
        if content_type in DOCX_TYPES:
            return _scan_text_based(
                client,
                _extract_text_from_docx(file_bytes),
                filename,
                content_type,
                model_id,
            )
        if content_type == "text/plain":
            return _scan_text_based(
                client,
                file_bytes.decode("utf-8", errors="replace"),
                filename,
                content_type,
                model_id,
            )
        # Fallback: try DOCX extraction (legacy behavior)
        return _scan_text_based(
            client,
            _extract_text_from_docx(file_bytes),
            filename,
            content_type,
            model_id,
        )
    except Exception:
        logger.exception("Document scan failed for %s", filename)
        return {
            "status": FileStatus.SCAN_ERROR,
            "scan_result": "Errore durante la scansione del documento.",
            "compiled": False,
            "confidence": "low",
            "cost_usd": 0.0,
            "tokens": {"input": 0, "output": 0, "total": 0},
        }


def _scan_pdf(
    client: anthropic.Anthropic,
    file_bytes: bytes,
    filename: str,
    content_type: str,
    model_id: str,
) -> dict[str, Any]:
    """Scan a PDF using Claude's document understanding (base64 input)."""
    b64_data = base64.b64encode(file_bytes).decode("utf-8")

    message = client.messages.create(
        model=model_id,
        max_tokens=512,
        system=SCAN_SYSTEM_PROMPT,
        messages=[
            {
                "role": "user",
                "content": [
                    {
                        "type": "document",
                        "source": {
                            "type": "base64",
                            "media_type": "application/pdf",
                            "data": b64_data,
                        },
                    },
                    {
                        "type": "text",
                        "text": (
                            f"Analizza questo documento PDF.\n"
                            f"Nome file: {filename}\n"
                            f"Determina se e' stato compilato o e' ancora un template vuoto."
                        ),
                    },
                ],
            }
        ],
        tools=[{"name": SCAN_TOOL_NAME, "description": SCAN_TOOL_DESCRIPTION, "input_schema": SCAN_INPUT_SCHEMA}],
        tool_choice={"type": "tool", "name": SCAN_TOOL_NAME},
    )

    return _parse_scan_response(message, model_id)


def _scan_text_based(
    client: anthropic.Anthropic,
    text_content: str,
    filename: str,
    content_type: str,
    model_id: str,
) -> dict[str, Any]:
    """Scan a text-based document (DOCX, XLSX, TXT) by sending extracted text to Claude."""
    if not text_content.strip():
        return {
            "status": FileStatus.NOT_COMPILED,
            "scan_result": "Il documento e' vuoto (nessun testo trovato).",
            "compiled": False,
            "confidence": "high",
            "cost_usd": 0.0,
            "tokens": {"input": 0, "output": 0, "total": 0},
        }

    # Truncate to avoid huge API costs
    truncated = text_content[:8000]

    user_prompt = SCAN_USER_PROMPT.format(
        filename=filename,
        content_type=content_type,
        content=truncated,
    )

    message = client.messages.create(
        model=model_id,
        max_tokens=512,
        system=SCAN_SYSTEM_PROMPT,
        messages=[{"role": "user", "content": user_prompt}],
        tools=[{"name": SCAN_TOOL_NAME, "description": SCAN_TOOL_DESCRIPTION, "input_schema": SCAN_INPUT_SCHEMA}],
        tool_choice={"type": "tool", "name": SCAN_TOOL_NAME},
    )

    return _parse_scan_response(message, model_id)


def _parse_scan_response(message: anthropic.types.Message, model_id: str) -> dict[str, Any]:
    """Map the tool_use payload from Claude to the FileStatus-shaped dict."""
    cost = _calculate_cost(message.usage, model_id)
    tokens = {
        "input": message.usage.input_tokens,
        "output": message.usage.output_tokens,
        "total": message.usage.input_tokens + message.usage.output_tokens,
    }

    tool_input: dict[str, Any] | None = None
    for block in message.content:
        if getattr(block, "type", None) == "tool_use":
            tool_input = dict(block.input) if isinstance(block.input, dict) else None
            break

    if tool_input is None:
        logger.warning("No tool_use block in scan response (model=%s)", model_id)
        return {
            "status": FileStatus.SCAN_ERROR,
            "scan_result": "Errore nel parsing della risposta AI.",
            "compiled": False,
            "confidence": "low",
            "cost_usd": cost,
            "tokens": tokens,
        }

    compiled = bool(tool_input.get("compiled", False))
    confidence = str(tool_input.get("confidence", "low"))
    summary = str(tool_input.get("summary", ""))

    return {
        "status": FileStatus.COMPILED if compiled else FileStatus.NOT_COMPILED,
        "scan_result": summary,
        "compiled": compiled,
        "confidence": confidence,
        "cost_usd": cost,
        "tokens": tokens,
    }
