"""Cover letter service."""

import io
import re
from datetime import datetime, timezone
from uuid import UUID

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt
from sqlalchemy.orm import Session

from ..analysis.models import JobAnalysis
from ..integrations.anthropic_client import generate_cover_letter
from ..integrations.cache import CacheService
from .models import CoverLetter


def create_cover_letter(
    db: Session,
    analysis: JobAnalysis,
    cv_text: str,
    language: str,
    model: str = "haiku",
    cache: CacheService | None = None,
) -> tuple[CoverLetter, dict]:
    """Generate and persist a cover letter for an analysis."""
    analysis_data = {
        "role": analysis.role,
        "company": analysis.company,
        "score": analysis.score,
        "strengths": analysis.strengths or [],
        "gaps": analysis.gaps or [],
    }

    result = generate_cover_letter(
        cv_text, analysis.job_description, analysis_data, language, model, cache
    )

    cl = CoverLetter(
        analysis_id=analysis.id,
        language=language,
        content=result.get("cover_letter", ""),
        subject_lines=result.get("subject_lines", []),
        model_used=result.get("model_used", ""),
        tokens_input=result.get("tokens", {}).get("input", 0),
        tokens_output=result.get("tokens", {}).get("output", 0),
        cost_usd=result.get("cost_usd", 0.0),
    )
    db.add(cl)
    db.flush()
    return cl, result


def get_cover_letter_by_id(db: Session, cover_letter_id: str) -> CoverLetter | None:
    """Fetch a cover letter by its UUID."""
    try:
        uid = UUID(cover_letter_id)
    except (ValueError, AttributeError):
        return None
    return db.query(CoverLetter).filter(CoverLetter.id == uid).first()


def build_docx(cover_letter: CoverLetter, analysis: JobAnalysis) -> io.BytesIO:
    """Generate a professionally formatted DOCX from a cover letter.

    Returns an in-memory BytesIO buffer ready to be sent as a response.
    """
    doc = Document()

    # -- Page margins: 2.5 cm all sides --
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # -- Default font: Calibri 11pt --
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)
    pf = style.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(6)
    pf.line_spacing = 1.15

    # -- Header: name, email, phone (right-aligned) --
    header_para = doc.add_paragraph()
    header_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    name_run = header_para.add_run("Marco Bellingeri\n")
    name_run.bold = True
    name_run.font.size = Pt(13)
    name_run.font.name = "Calibri"
    contact_run = header_para.add_run("marco.bellingeri@gmail.com | +39 348 450 7859")
    contact_run.font.size = Pt(9)
    contact_run.font.name = "Calibri"
    contact_run.font.color.rgb = None  # inherit (dark gray from default)

    # -- Date --
    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    now = datetime.now(timezone.utc)
    date_str = now.strftime("%d/%m/%Y")
    date_run = date_para.add_run(date_str)
    date_run.font.size = Pt(10)
    date_run.font.name = "Calibri"

    # -- Spacer --
    doc.add_paragraph()

    # -- Letter body: split by double newlines into paragraphs --
    content = cover_letter.content or ""
    # Normalize literal \n sequences (from JSON) into actual newlines
    content = content.replace("\\n", "\n")
    paragraphs = re.split(r"\n{2,}", content.strip())

    for para_text in paragraphs:
        # Clean up single newlines within a paragraph (join lines)
        cleaned = para_text.strip().replace("\n", " ")
        if not cleaned:
            continue
        p = doc.add_paragraph(cleaned)
        p.paragraph_format.space_after = Pt(8)
        for run in p.runs:
            run.font.name = "Calibri"
            run.font.size = Pt(11)

    # -- Build filename --
    company = (analysis.company or "").strip()
    if company:
        # Sanitize company name for filename
        safe_company = re.sub(r'[<>:"/\\|?*]', "", company)
        safe_company = re.sub(r"\s+", "_", safe_company)
        filename = f"Cover_Letter_{safe_company}.docx"
    else:
        filename = "Cover_Letter.docx"

    # -- Write to buffer --
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)

    return buf, filename
